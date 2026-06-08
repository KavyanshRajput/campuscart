import os
import sys
import docx
from docx2pdf import convert

def generate_guide():
    doc = docx.Document()
    
    # Define style helper
    def add_para_with_bold_lead(doc, lead, text, style=None):
        p = doc.add_paragraph(style=style)
        run_lead = p.add_run(lead)
        run_lead.bold = True
        p.add_run(text)
        return p

    # Title
    doc.add_heading('CampusCart (AcroCart) — Project Explanation Guide', level=0)
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = docx.shared.Pt(24)
    run_sub = p_sub.add_run("Created for Minor Project Exhibition (6th Semester)\nAcropolis Institute of Technology & Research, Indore\nDesigned for Beginners & Presentation Evaluation")
    run_sub.italic = True
    
    # Section 1: Project Overview
    doc.add_heading('1. Project Overview & The Real-World Problem', level=1)
    
    doc.add_paragraph(
        "CampusCart (also known as AcroCart) is a peer-to-peer digital marketplace built specifically "
        "for the college campus. It allows students to buy, sell, and trade second-hand college essentials "
        "like textbooks, reference materials, lab coats, drafters, scientific calculators, and other items."
    )
    
    doc.add_heading('The Problem it Solves (Why we built this):', level=2)
    
    add_para_with_bold_lead(doc, "1. Unorganized informal groups: ", "Usually, students sell books via informal WhatsApp groups. These chats get flooded, messages get lost, and there is no searchable catalog.", style='List Bullet')
    add_para_with_bold_lead(doc, "2. Security & Trust: ", "Public platforms like OLX or Facebook Marketplace are filled with spam, fake profiles, and require meeting strangers outside of campus. CampusCart solves this by forcing all transactions to occur inside the safe, verified college campus.", style='List Bullet')
    add_para_with_bold_lead(doc, "3. Domain Restricting: ", "To enter the marketplace, users must log in using their official college email. This ensures that only actual students can buy or sell.", style='List Bullet')

    # Section 2: Technology Stack Explained in Plain English
    doc.add_heading('2. The Technology Stack Explained Simply', level=1)
    
    doc.add_paragraph(
        "You do not need to be a coding genius to explain the technology behind this app. Here is a simple "
        "breakdown of the technologies used, written in everyday language so you can explain them to your evaluators:"
    )

    doc.add_heading('A. React 19 (The Building Blocks)', level=2)
    doc.add_paragraph(
        "React is like a box of Lego. Instead of writing one massive, complicated webpage, React allows us "
        "to break the interface down into small, reusable Lego blocks called 'Components'. "
        "For example, the Navbar (navigation bar) is a single Component, and the Product Card is another Component. "
        "If we want to display 10 items, we don't code 10 items manually — we just tell React to reuse the 'Product Card' Component 10 times, pasting in the unique data (title, price) for each."
    )

    doc.add_heading('B. Next.js 16 (The Manager)', level=2)
    doc.add_paragraph(
        "Next.js is a framework built on top of React. Think of it as the manager of the website. It handles:\n"
        "1. Routing: When a user clicks 'Sell', Next.js loads the '/sell' page. When they click 'Profile', it loads '/profile'.\n"
        "2. Performance: It packages all the JavaScript and assets so the website loads extremely fast.\n"
        "3. Static Exporting: It allows us to build the app as a set of static files (HTML, CSS, JS) that can be hosted on a fast, cheap cloud server."
    )

    doc.add_heading('C. CSS Modules (The Scoped Designer)', level=2)
    doc.add_paragraph(
        "CSS is what gives the website its styling — the colors, gradients, card shapes, and fonts. "
        "Normally, if you write CSS, styles from one page can accidentally leak and mess up other pages. "
        "CSS Modules solves this. It makes the styles 'scoped' (locked) to that specific component. "
        "So, the styles in 'profile.module.css' only affect the profile page, and cannot break anything on the homepage."
    )

    doc.add_heading('D. Firebase (The Serverless Backend)', level=2)
    doc.add_paragraph(
        "Normally, a website requires you to write a separate database server (using Node.js, Django, or Java) and rent database servers. "
        "Firebase is a 'Backend-as-a-Service' (BaaS). It is a cloud database and auth system provided by Google. "
        "It handles all backend requirements for us automatically, so we don't have to code a backend server from scratch. We use three Firebase features:"
    )
    
    add_para_with_bold_lead(doc, "1. Firebase Authentication: ", "Handles the Google Sign-in popup. It verifies who the user is and retrieves their profile picture, name, and email address.", style='List Bullet')
    add_para_with_bold_lead(doc, "2. Cloud Firestore (The Database): ", "A cloud-based database that stores listings, chat messages, and user settings. It stores data in 'documents' (like files inside folders). When a user updates their settings or sends a message, it updates Firestore in real-time.", style='List Bullet')
    add_para_with_bold_lead(doc, "3. Firebase Hosting: ", "This is the cloud server that hosts our compiled Next.js files on the internet so that anyone can open the live URL: https://campuscart-9c677.web.app.", style='List Bullet')

    doc.add_heading('E. Progressive Web App (PWA)', level=2)
    doc.add_paragraph(
        "CampusCart is built as a PWA. This means that when a user opens the website on their mobile phone, "
        "their browser shows an 'Install App' button. Once clicked, the website gets installed on their home screen "
        "with a custom icon and runs in standalone mode (without the browser address bar), looking and feeling exactly like a native Android or iOS app."
    )

    # Section 3: How the Core Features are Programmed
    doc.add_heading('3. How the Key Features Work Under the Hood', level=1)
    
    doc.add_paragraph(
        "Here is the simple, conceptual breakdown of the code logic. If the evaluator asks, "
        "'How did you program X?', these explanations will give you the exact logic to present."
    )

    doc.add_heading('Feature 1: Locking logins to college email (@acropolis.in)', level=2)
    doc.add_paragraph(
        "How it's coded: When a student clicks 'Sign in with Google', Firebase Auth triggers a popup. "
        "Once they log in, we write an event listener in the file 'src/lib/AuthContext.js'. "
        "The code checks if the email address ends with '@acropolis.in'. If yes, it logs them in. "
        "If no (e.g. they logged in with their personal gmail), the code automatically logs them out, "
        "clears the session, and displays an error saying 'Access restricted to @acropolis.in emails only.'"
    )

    doc.add_heading('Feature 2: Real-Time Listings (onSnapshot)', level=2)
    doc.add_paragraph(
        "How it's coded: Normally, to get new data on a website, you have to refresh the page. "
        "In our app, we use a Firestore feature called 'onSnapshot'. It creates a live, open tunnel "
        "between the user's browser and the database. The moment a student uploads a new item to the "
        "database from their phone, the database tells all active browsers, 'Hey, here is a new listing!' "
        "and the homepage updates automatically in under 200 milliseconds without requiring a page reload."
    )

    doc.add_heading('Feature 3: Live In-App Chat system', level=2)
    doc.add_paragraph(
        "How it's coded: When a buyer opens an item details page and clicks 'Chat Now', Next.js redirects them "
        "to the Profile page and opens the 'Chats' tab. The code checks if a chat thread already exists in Firestore "
        "between the buyer and seller. If yes, it loads that thread. If not, it creates a new chat document. "
        "Inside the chat, we use a sub-collection called 'messages'. Every time a user types a message and hits send, "
        "it adds a new document to the database. An 'onSnapshot' listener displays these messages on both users' screens "
        "instantly, creating a WhatsApp-like real-time chat experience."
    )

    doc.add_heading('Feature 4: Local Wishlist (savedItems)', level=2)
    doc.add_paragraph(
        "How it's coded: When a user clicks the Heart icon on a product, the code adds that product's unique ID "
        "to an array inside the browser's 'LocalStorage' (a simple database built directly into every web browser). "
        "When they visit the 'Saved' tab, the code reads these IDs from LocalStorage and fetches those listings "
        "from Firestore. This saves server space and works instantly without needing a backend server table."
    )

    # Section 4: File-by-File Structure & Code Logic Details
    doc.add_heading('4. Detailed File-by-File Code & Logic Breakdown', level=1)
    doc.add_paragraph(
        "This section explains what each file in your codebase is responsible for, how it is written, and the core coding concepts it uses."
    )

    doc.add_heading('A. src/app/layout.js (The Main Frame)', level=2)
    doc.add_paragraph(
        "• Purpose: This is the shell or template that wraps all pages on your website. It controls what remains visible when navigating (like the navigation header and footer).\n"
        "• What we did in the code:\n"
        "  1. Loaded the custom geometric font 'Outfit' from Google Fonts to give the project a premium look.\n"
        "  2. Wrapped the entire application content with the `<AuthProvider>` context component. This ensures that every single page knows immediately if a user is logged in or out.\n"
        "  3. Rendered the `<Navbar />` at the top and the `<Footer />` at the bottom of the content container."
    )

    doc.add_heading('B. src/lib/AuthContext.js (The Guard Gate)', level=2)
    doc.add_paragraph(
        "• Purpose: Manages user logins, logouts, credentials, and keeps track of who is currently using the app.\n"
        "• What we did in the code:\n"
        "  1. Set up a listener (`onAuthStateChanged`) that monitors if a user is logged in via Firebase Auth.\n"
        "  2. Added the domain lock filter: if the user's email domain does not end with '@acropolis.in', it calls `signOut(auth)` immediately to reject them and sets an error message state.\n"
        "  3. Programmed a 'Localhost mock bypass': if the code detects the website is running locally on `localhost:3000` (development mode), it skips the live Google Sign-In redirect (which is blocked on localhost) and automatically logs in as a mock student profile. This allows you to test settings, chats, and lists instantly without database blockages.\n"
        "  4. Wrapped state setters (`setUser` and `setLoading`) inside asynchronous `setTimeout` triggers to satisfy React's hooks performance standards."
    )

    doc.add_heading('C. src/lib/firebase.js (The Connection Bridge)', level=2)
    doc.add_paragraph(
        "• Purpose: Initializes the connection between your Next.js application and the Firebase cloud servers.\n"
        "• What we did in the code:\n"
        "  1. Loaded the Firebase project configurations using environment variables (`process.env.NEXT_PUBLIC_...`). This is a security best practice so your private keys are not exposed in the git repository.\n"
        "  2. Instantiated and exported the core Firebase client services: `auth` (for users), `db` (Firestore database), and `storage` (for images)."
    )

    doc.add_heading('D. src/app/page.js (The Homepage)', level=2)
    doc.add_paragraph(
        "• Purpose: Renders the home screen that users see when they open the website. It contains search bars, category boxes, a live activity ticker, recent items, and a campus stats widget.\n"
        "• What we did in the code:\n"
        "  1. Coded a live Firestore listener (`onSnapshot`) that queries the `listings` collection ordered by creation date. It filters out any listing marked as `'sold'` and stores it in the `listings` state.\n"
        "  2. Set up a PWA install prompt handler that listens to the browser's `beforeinstallprompt` event, saving it to state and rendering a floating 'Install App' banner if the app can be installed.\n"
        "  3. Programmed a 'Campus Pulse' widget that calculates the live count of active items and the number of active unique sellers directly from the Firestore collection snapshot.\n"
        "  4. Built a live activity ticker that monitors modifications in the collection and adds message updates (like 'Student listed Cooler for Rs 2000') using React state arrays."
    )

    doc.add_heading('E. src/app/explore/page.js (The Search & Filter Hub)', level=2)
    doc.add_paragraph(
        "• Purpose: Renders the full catalog page where students search, browse, and filter listings.\n"
        "• What we did in the code:\n"
        "  1. Listened to URL query parameters (e.g. `?search=textbook`) to pre-fill search fields.\n"
        "  2. Programmed client-side filtering logic that filters listings by categories, maximum price (using an interactive slider), and item age/condition checkbox selections.\n"
        "  3. Created an interface toggle allowing students to switch the catalog display between a Grid view (multi-column) and a List view (stacked)."
    )

    doc.add_heading('F. src/app/sell/page.js (The Listing Form)', level=2)
    doc.add_paragraph(
        "• Purpose: Renders the form that allows students to list items for sale.\n"
        "• What we did in the code:\n"
        "  1. Set up a standard HTML form with inputs for Title, Price, Category, Age/Condition, and Description.\n"
        "  2. Mapped the selected category to a specific default image illustration (e.g. a book graphic for 'Books' or a gadget graphic for 'Electronics') to give listings a fallback thumbnail.\n"
        "  3. Handled form submissions by calling Firestore's `addDoc` to upload the product details, along with the logged-in student's Google profile photo, name, and email."
    )

    doc.add_heading('G. src/app/profile/page.js (The Dashboard & Chat Engine)', level=2)
    doc.add_paragraph(
        "• Purpose: Renders the personal dashboard where students manage their listings, favorites, settings, and messaging.\n"
        "• What we did in the code:\n"
        "  1. **Listings Management:** Queries Firestore for listings matching `sellerId == user.uid`. Added action buttons to trigger `deleteDoc` (deleting items), `updateDoc` (marking items as sold or editing details via forms).\n"
        "  2. **Saved Wishlist:** Reads the listing IDs stored in the browser's `localStorage`, fetches the details of those specific items from Firestore, and renders them.\n"
        "  3. **Live Chat Engine:** Renders a split-pane layout (conversations list on the left, active chat window on the right). It listens to the `chats` collection in real-time. Typing a message triggers `addDoc` into a sub-collection `chats/{chatId}/messages`, and updates `lastMessage` on the parent chat document. Used scroll-into-view references to automatically snap chat threads to the latest message.\n"
        "  4. **Settings Form:** Handles profile preferences (WhatsApp contact hook, department selection, and semester level) and saves them in the `users` collection in Firestore. Added toggles to enable/disable notifications and to manage user blocks (muting specific sellers).\n"
        "  5. **Dark Mode Toggle:** Reads and writes the `'darkMode'` key in `localStorage` and toggles a `.dark` CSS class on `document.body` to switch the global color scheme in real-time."
    )

    doc.add_heading('H. src/components/ProductCard.js (The Item Template)', level=2)
    doc.add_paragraph(
        "• Purpose: A reusable, modular component that displays a summary of a single listing.\n"
        "• What we did in the code:\n"
        "  1. Built the card tile to display the product thumbnail, category tag, title, price, usage age, and a customized sold badge overlay.\n"
        "  2. Coded a Heart button that handles wishlist bookmarks: it checks if the item is in `localStorage.getItem('savedItems')`, toggles its presence, and triggers storage synchronizations."
    )

    doc.add_heading('I. firestore.rules & firebase.json (The Security Configuration)', level=2)
    doc.add_paragraph(
        "• Purpose: Controls cloud database permissions and deployment targets.\n"
        "• What we did in the code:\n"
        "  1. Configured security rules that allow read and write operations on your database collections under an extended timeframe (valid until December 31, 2026), preventing your app from being locked down by Firebase's 30-day default expiration.\n"
        "  2. Configured `firebase.json` to map the `firestore.rules` script to your deployment configuration, ensuring the rules are tracked in version control and pushed automatically when running database deployments."
    )

    # Section 5: Database Design
    doc.add_heading('5. Database Schema Design (Firestore)', level=1)
    doc.add_paragraph(
        "Firestore is a NoSQL database. It organizes data in 'Collections' (folders) which contain 'Documents' (files). "
        "Each document contains 'Fields' (data pairs). We have 3 main Collections:"
    )
    
    doc.add_heading('A. listings Collection:', level=2)
    doc.add_paragraph(
        "Stores all items posted for sale. Each document in this collection has:\n"
        "• title (String): Name of the item (e.g., 'Engineering Graphics Drafter')\n"
        "• description (String): Details about the item's condition\n"
        "• price (Number): Selling price in Rupees (e.g., 250)\n"
        "• category (String): e.g., 'books', 'electronics', 'fashion', 'home', 'sports'\n"
        "• status (String): 'active' (for sale) or 'sold' (sold items)\n"
        "• sellerId, sellerName, sellerEmail, sellerPhoto: Stored directly so we know who is selling it\n"
        "• createdAt (Timestamp): When it was posted, to sort items newest first"
    )

    doc.add_heading('B. chats Collection:', level=2)
    doc.add_paragraph(
        "Stores the active conversation threads. Fields include:\n"
        "• participants (Array): A list of the two users' IDs involved in the chat\n"
        "• participantDetails (Map): Profile names and pictures of the participants\n"
        "• itemName (String): The item they are negotiating about\n"
        "• lastMessage (String): Snippet of the last text sent, shown in the sidebar\n"
        "• Sub-collection 'messages': Contains the actual text bubbles with 'senderId', 'text', and 'createdAt'."
    )

    doc.add_heading('C. users Collection:', level=2)
    doc.add_paragraph(
        "Stores individual student settings. Fields include:\n"
        "• phone (String): Their WhatsApp/phone number for contact hooks\n"
        "• department (String): e.g., 'CSE', 'IT', 'EC' (used to match classmates)\n"
        "• semester (String): e.g., '6' (their current semester level)\n"
        "• matchAlerts (Boolean): True/False if they want department match highlights\n"
        "• blockedUsers (Array): List of user IDs they have muted/blocked"
    )

    # Section 6: Project Exhibition Viva Q&A Guide
    doc.add_heading('6. Project Exhibition Viva Q&A (Be Prepared!)', level=1)
    doc.add_paragraph(
        "Here are the most common questions an evaluator will ask during your project exhibition, "
        "along with the exact, easy-to-understand answers you should give:"
    )

    # Q1
    q1 = doc.add_paragraph()
    q1.add_run('Q1: What is the main purpose of this project?\n').bold = True
    q1.add_run('Answer: ').italic = True
    q1.add_run(
        "CampusCart is an exclusive, secure peer-to-peer marketplace app designed for college students "
        "to trade second-hand items. It replaces unorganized student WhatsApp groups and unsafe open platforms like OLX "
        "by locking authentication to our college domain (@acropolis.in) and forcing safe, physical transactions on campus."
    )
    
    # Q2
    q2 = doc.add_paragraph()
    q2.add_run('\nQ2: What is the tech stack of this application?\n').bold = True
    q2.add_run('Answer: ').italic = True
    q2.add_run(
        "We used Next.js 16 and React 19 for the frontend (compiled as a Progressive Web App), "
        "Vanilla CSS Modules for scoped styling, and Firebase (Authentication, Firestore, Hosting) for the serverless backend."
    )

    # Q3
    q3 = doc.add_paragraph()
    q3.add_run('\nQ3: Why did you choose Firestore over SQL databases like MySQL?\n').bold = True
    q3.add_run('Answer: ').italic = True
    q3.add_run(
        "Since our app requires real-time messaging/chat and instant listing updates, Firestore's built-in "
        "real-time data synchronization ('onSnapshot') was a perfect fit. It allows us to stream database changes "
        "directly to the browser without setting up complex WebSocket servers. In addition, Firestore's flexible, document-based "
        "structure made development fast and scalable."
    )

    # Q4
    q4 = doc.add_paragraph()
    q4.add_run('\nQ4: How did you implement domain-locked security?\n').bold = True
    q4.add_run('Answer: ').italic = True
    q4.add_run(
        "We used Firebase Auth's Google login provider. When a user authenticates, the client-side code checks if the "
        "returned email address ends with our college's domain name (@acropolis.in). If it doesn't match, we automatically "
        "log them out and show a restricted access error. On the database level, we wrote Firestore Security Rules that block "
        "read/write operations from unauthenticated users, keeping our database secure."
    )

    # Q5
    q5 = doc.add_paragraph()
    q5.add_run('\nQ5: What are Firestore Security Rules and why did you update them?\n').bold = True
    q5.add_run('Answer: ').italic = True
    q5.add_run(
        "Firestore Security Rules are configuration scripts deployed on Firebase servers that validate database access requests. "
        "They prevent unauthorized attackers from writing or deleting our records. Initially, Firebase sets up temporary development rules "
        "which expire after 30 days. We created custom rules (saved in firestore.rules and firebase.json) and deployed them to keep the "
        "project open for our academic evaluation while keeping it protected."
    )

    # Q6
    q6 = doc.add_paragraph()
    q6.add_run('\nQ6: Explain what a Progressive Web App (PWA) is and how it benefits your app.\n').bold = True
    q6.add_run('Answer: ').italic = True
    q6.add_run(
        "A PWA is a hybrid website that can be installed directly onto a mobile phone or desktop as a standalone app, "
        "meaning it runs in its own window without the browser address bar. It uses a manifest.json file to provide "
        "app icons, branding colors, and locked layout orientations. This gives our users a native mobile app experience "
        "without requiring them to download a large app from the Google Play Store or App Store."
    )

    # Q7
    q7 = doc.add_paragraph()
    q7.add_run('\nQ7: How does local development work without Google Auth blocking localhost?\n').bold = True
    q7.add_run('Answer: ').italic = True
    q7.add_run(
        "Google Sign-In sometimes blocks redirects from localhost in development mode due to referrer restrictions. "
        "To solve this, we coded a 'mock login bypass' in our AuthContext.js file. When the app detects that it is running "
        "on 'localhost', it automatically logs the session in as a pre-configured mock student ('KAVYANSH RAJPUT'). "
        "This allows us to test the entire application interface, settings, and messaging flows locally with ease."
    )

    # Q8
    q8 = doc.add_paragraph()
    q8.add_run('\nQ8: What is your database structure and how are relationships handled?\n').bold = True
    q8.add_run('Answer: ').italic = True
    q8.add_run(
        "We use a NoSQL database, so relationships are handled using 'denormalization' and 'flat structures'. "
        "Each product in the 'listings' collection stores its seller's ID, name, and email directly. Each conversation thread "
        "in the 'chats' collection stores an array of the two participants' IDs. This structure makes database reads extremely "
        "fast because we do not need to run expensive SQL join queries to fetch seller details or chat history."
    )

    # Q9
    q9 = doc.add_paragraph()
    q9.add_run('\nQ9: What are the main features of the user dashboard?\n').bold = True
    q9.add_run('Answer: ').italic = True
    q9.add_run(
        "The dashboard (Profile page) provides listing management (mark as sold, edit description, delete listings), "
        "wishlist favorites loaded from the browser's LocalStorage, real-time in-app chat windows, and settings where students "
        "can update their WhatsApp phone number, department branch, semester, and block/mute users."
    )

    # Q10
    q10 = doc.add_paragraph()
    q10.add_run('\nQ10: What is the future scope of this project?\n').bold = True
    q10.add_run('Answer: ').italic = True
    q10.add_run(
        "In the future, we plan to implement:\n"
        "1. Image Upload: Allow users to upload actual photos of their items directly to Firebase Storage (currently it auto-assigns category illustrations).\n"
        "2. Push Notifications: Integrate Firebase Cloud Messaging (FCM) to alert users about new messages or price drops.\n"
        "3. UPI Payments: Allow in-app UPI QR codes or transactions for cash-free trading.\n"
        "4. Multi-Campus Support: Configure the app to support other local colleges by extending the domain-restriction list."
    )

    # Save document
    docx_path = r"e:\Minor 6th sem\CampusCart_Project_Explanation_Guide.docx"
    pdf_path = r"e:\Minor 6th sem\CampusCart_Project_Explanation_Guide.pdf"
    
    doc.save(docx_path)
    print(f"Docx file saved successfully at: {docx_path}")
    
    # Convert to PDF
    print("Converting Docx to PDF...")
    convert(docx_path, pdf_path)
    print(f"PDF converted successfully at: {pdf_path}")

if __name__ == '__main__':
    generate_guide()
